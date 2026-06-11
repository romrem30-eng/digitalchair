from io import BytesIO
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph
from reportlab.platypus import SimpleDocTemplate
from reportlab.platypus import Spacer
from reportlab.platypus import Table
from reportlab.platypus import TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


PDF_FONT_NAME = 'Helvetica'


def get_pdf_font_name():

    font_candidates = (
        Path('C:/Windows/Fonts/arial.ttf'),
        Path('C:/Windows/Fonts/calibri.ttf'),
        Path('C:/Windows/Fonts/tahoma.ttf'),
    )

    for font_path in font_candidates:

        if font_path.exists():

            try:

                pdfmetrics.registerFont(
                    TTFont(
                        'DigitalChairUnicode',
                        str(font_path)
                    )
                )

                return 'DigitalChairUnicode'

            except Exception:

                continue

    return PDF_FONT_NAME


def autosize_columns(worksheet):

    for column_cells in worksheet.columns:

        max_length = 0
        column_letter = column_cells[0].column_letter

        for cell in column_cells:

            cell_value = '' if cell.value is None else str(cell.value)
            max_length = max(
                max_length,
                len(cell_value)
            )

        worksheet.column_dimensions[column_letter].width = min(
            max_length + 2,
            40
        )


def workbook_to_bytes(workbook):

    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def build_workload_summary_xlsx(assignments):

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = 'Workload Summary'

    worksheet.append([
        'Преподаватель',
        'Дисциплина',
        'Часы',
        'Семестр',
        'Учебный год',
    ])

    for assignment in assignments:

        worksheet.append([
            assignment.teacher.full_name,
            assignment.subject.name,
            assignment.assigned_hours,
            assignment.semester,
            assignment.academic_year,
        ])

    autosize_columns(worksheet)
    return workbook_to_bytes(workbook)


def build_workload_execution_xlsx(plan_rows):

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = 'Workload Execution'

    worksheet.append([
        'Кафедра',
        'Учебный год',
        'Плановые часы',
        'Распределенные часы',
        'Оставшиеся часы',
        'Статус плана',
    ])

    for row in plan_rows:

        worksheet.append([
            row['plan'].кафедра,
            row['plan'].academic_year,
            row['plan'].total_hours,
            row['distributed_hours'],
            row['remaining_hours'],
            row['plan'].get_status_display(),
        ])

    autosize_columns(worksheet)
    return workbook_to_bytes(workbook)


def build_tasks_report_xlsx(tasks):

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = 'Tasks Report'

    worksheet.append([
        'Название поручения',
        'Исполнитель',
        'Срок',
        'Статус',
        'Приоритет',
        'Дата завершения',
    ])

    for task in tasks:

        worksheet.append([
            task.title,
            task.teacher.full_name,
            str(task.due_date),
            task.get_status_display(),
            task.get_priority_display(),
            task.completed_at.strftime('%Y-%m-%d %H:%M') if task.completed_at else '-',
        ])

    autosize_columns(worksheet)
    return workbook_to_bytes(workbook)


def build_contingent_xlsx(teachers_count, subjects_count, groups_count):

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = 'Contingent'

    worksheet.append([
        'Показатель',
        'Количество',
    ])

    worksheet.append(['Преподаватели', teachers_count])
    worksheet.append(['Дисциплины', subjects_count])
    worksheet.append(['Группы', groups_count])

    autosize_columns(worksheet)
    return workbook_to_bytes(workbook)


def build_teacher_workload_docx(selected_teacher, assignments, total_hours):

    document = Document()
    document.add_heading('Индивидуальная нагрузка преподавателя', level=1)

    if selected_teacher is not None:

        document.add_paragraph(
            f'Преподаватель: {selected_teacher.full_name}'
        )

        document.add_paragraph(
            'Должность: '
            f'{selected_teacher.get_position_display()}'
        )

    document.add_paragraph(f'Итоговая нагрузка: {total_hours} ч.')

    table = document.add_table(
        rows=1,
        cols=4
    )
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Дисциплина'
    header_cells[1].text = 'Часы'
    header_cells[2].text = 'Семестр'
    header_cells[3].text = 'Учебный год'

    for assignment in assignments:

        row_cells = table.add_row().cells
        row_cells[0].text = assignment.subject.name
        row_cells[1].text = str(assignment.assigned_hours)
        row_cells[2].text = str(assignment.semester)
        row_cells[3].text = assignment.academic_year

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_teacher_workload_pdf(selected_teacher, assignments, total_hours):

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=15 * mm,
        leftMargin=15 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )

    styles = getSampleStyleSheet()
    font_name = get_pdf_font_name()

    styles['Title'].fontName = font_name
    styles['Normal'].fontName = font_name

    story = [
        Paragraph(
            'Индивидуальная нагрузка преподавателя',
            styles['Title']
        ),
        Spacer(1, 8),
    ]

    if selected_teacher is not None:

        story.append(
            Paragraph(
                f'Преподаватель: {selected_teacher.full_name}',
                styles['Normal']
            )
        )
        story.append(
            Paragraph(
                'Должность: '
                f'{selected_teacher.get_position_display()}',
                styles['Normal']
            )
        )

    story.append(
        Paragraph(
            f'Итоговая нагрузка: {total_hours} ч.',
            styles['Normal']
        )
    )
    story.append(Spacer(1, 10))

    table_data = [[
        'Дисциплина',
        'Часы',
        'Семестр',
        'Учебный год',
    ]]

    for assignment in assignments:

        table_data.append([
            assignment.subject.name,
            str(assignment.assigned_hours),
            str(assignment.semester),
            assignment.academic_year,
        ])

    table = Table(
        table_data,
        repeatRows=1
    )
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#dbeafe')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTNAME', (0, 0), (-1, 0), font_name),
        ('FONTNAME', (0, 1), (-1, -1), font_name),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('PADDING', (0, 0), (-1, -1), 6),
    ]))

    story.append(table)
    document.build(story)
    return buffer.getvalue()
